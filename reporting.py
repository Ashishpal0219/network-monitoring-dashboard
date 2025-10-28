import pandas as pd
from docx import Document
from docx.shared import Inches
import io

def export_to_csv(df):
    """
    Converts a DataFrame to a CSV string in memory.
    
    Returns:
        bytes: The CSV data, encoded in UTF-8.
    """
    # Reset index if 'timestamp' is the index
    if isinstance(df.index, pd.DatetimeIndex):
        df_export = df.reset_index()
    else:
        df_export = df.copy()
        
    return df_export.to_csv(index=False).encode('utf-8')

def export_to_docx(df, title):
    """
    Converts a DataFrame to a .docx file in memory.
    
    Args:
        df (pd.DataFrame): The data to export.
        title (str): The title to put at the top of the document.
        
    Returns:
        io.BytesIO: A_bytes buffer containing the .docx file data.
    """
    # Reset index if 'timestamp' is the index
    if isinstance(df.index, pd.DatetimeIndex):
        df_export = df.reset_index()
    else:
        df_export = df.copy()

    try:
        doc = Document()
        doc.add_heading(title, level=1)
        doc.add_paragraph(f"Report generated on {pd.to_datetime('now').strftime('%Y-%m-%d %H:%M:%S')}")

        # Add the table
        if not df_export.empty:
            table = doc.add_table(rows=1, cols=len(df_export.columns))
            table.style = 'Table Grid'
            
            # Add header row
            hdr_cells = table.rows[0].cells
            for i, col_name in enumerate(df_export.columns):
                hdr_cells[i].text = str(col_name)

            # Add data rows
            for _, row in df_export.iterrows():
                row_cells = table.add_row().cells
                for i, val in enumerate(row):
                    row_cells[i].text = str(val)

        # Save to a bytes buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    except Exception as e:
        print(f"Error creating DOCX file: {e}")
        return io.BytesIO() # Return empty buffer on error